# -*- coding: utf-8 -*-
"""将 发票识别技术方案_专利申请.md 转为 Word 文档。运行：python docs/md_to_docx.py（项目根目录）"""
import os
import re
from pathlib import Path

DOCS_DIR = Path(__file__).resolve().parent
IMAGES_DIR = DOCS_DIR / "images"
MD_PATH = DOCS_DIR / "发票识别技术方案_专利申请.md"
DOCX_PATH = DOCS_DIR / "发票识别技术方案_专利申请.docx"


def main():
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("请先安装: pip install python-docx")
        return 1

    if not MD_PATH.exists():
        print(f"未找到: {MD_PATH}")
        return 1

    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    lines = content.split("\n")
    i = 0
    in_mermaid = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```mermaid"):
            in_mermaid = True
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            doc.add_paragraph("（流程图见附图或 Mermaid 源码。）")
            doc.add_paragraph()
            continue

        if stripped.startswith("```") and "mermaid" not in stripped.lower():
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                i += 1
            if i < len(lines):
                i += 1
            continue

        if re.match(r"^# [^#]", line) and not line.startswith("## "):
            doc.add_heading(line.lstrip("# ").strip(), level=0)
            i += 1
            continue
        if line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line.lstrip("# ").strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line.lstrip("# ").strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line.lstrip("# ").strip(), level=3)
            i += 1
            continue

        img_match = re.match(r'^!\[([^\]]*)\]\((.*)\)\s*$', stripped)
        if img_match:
            alt, path = img_match.group(1), img_match.group(2)
            if "images" in path:
                img_name = os.path.basename(path)
                img_full = IMAGES_DIR / img_name
                if img_full.exists():
                    try:
                        doc.add_paragraph(alt or "附图")
                        doc.add_picture(str(img_full), width=Inches(5.5))
                        doc.add_paragraph()
                    except Exception as e:
                        doc.add_paragraph(f"（附图：{img_name}，未嵌入：{e}）")
                else:
                    doc.add_paragraph(f"（附图：{img_name}）")
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        para_lines = []
        while i < len(lines):
            ln = lines[i]
            if not ln.strip():
                break
            if (ln.startswith("# ") or ln.startswith("## ") or ln.startswith("### ") or
                    ln.startswith("#### ") or re.match(r'^!\[.*\]\(.*\)', ln.strip()) or
                    ln.strip().startswith("```")):
                break
            para_lines.append(ln)
            i += 1
        if para_lines:
            text = "\n".join(para_lines)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_paragraph(text)
            continue
        i += 1

    doc.save(DOCX_PATH)
    print(f"已生成: {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    exit(main())
